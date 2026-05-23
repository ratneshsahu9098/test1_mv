from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('MV Tax Reminder — Complete Project Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Prepared on: ').bold = True
p.add_run('23 May 2026')

doc.add_paragraph('─' * 80)

# ============================================================
# 1. OVERVIEW
# ============================================================
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'MV Tax Reminder is a full-stack Software-as-a-Service (SaaS) platform designed to automate '
    'Motor Vehicle (MV) tax renewal reminders for vehicle owners in India. The system scrapes '
    'the Indian government Parivahan portal to fetch vehicle tax expiry dates, provides a '
    'dashboard for fleet managers to monitor tax status, and sends multi-channel reminders '
    '(Email, Telegram, Browser Push, WhatsApp) at configurable intervals (7/3/1 days before '
    'expiry, plus after expiry). It includes a subscription billing system via Razorpay for '
    'monetization, role-based access control (Admin / Staff / Viewer), and is fully '
    'mobile-responsive with dark/light theme support.'
)

doc.add_heading('Tech Stack', level=2)
items = [
    ('Frontend', 'React 19, Vite 8, Tailwind CSS v3, framer-motion, recharts, lucide-react, axios, '
                 'Firebase Cloud Messaging (FCM) for browser push, Razorpay Web SDK'),
    ('Backend', 'Python 3.12, Flask, Flask-JWT-Extended (7-day tokens), Flask-CORS, APScheduler, '
                'SQLite 3, Bcrypt, Playwright (web scraping), Firebase Admin SDK, ReportLab (PDF), '
                'OpenPyXL (Excel export)'),
    ('Infrastructure', 'SQLite with WAL mode, 7 database indexes, gzip compression (~154 kB total frontend), '
                       'React.lazy code-splitting (18 chunks), connection pooling via get_db() helper'),
    ('External APIs', 'Razorpay (subscription payments), Parivahan.gov.in (vehicle data), '
                      'Firebase Cloud Messaging (push), UltraMSG (WhatsApp), '
                      'Gmail SMTP (email), Telegram Bot API'),
]
for label, value in items:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_paragraph('')

# ============================================================
# 2. DIRECTORY STRUCTURE
# ============================================================
doc.add_heading('2. Directory Structure', level=1)

structure = """D:\\Vehicle Reminder\\
├── dashboard.py              # Flask app entry point, DB init, dashboard API, legacy HTML routes
├── fetch_vehicle.py          # Parivahan web scraper using Playwright (Chromium)
├── helpers.py                # Shared utilities: DB_FILE, get_current_user, auto_downgrade, WhatsApp
├── run.py                    # Alternative standalone entry point
├── requirements.txt          # Python package dependencies
├── Procfile                  # Heroku deployment configuration
├── .env                      # Environment secrets (Telegram, Gmail SMTP, Firebase, CORS)
├── firebase-service-account.json  # Firebase Admin SDK credentials (service account key)
├── vehicles.db               # SQLite database (auto-created)
│
├── config/
│   └── settings.py           # Centralized env variable config: Telegram, Email, FCM, reminder tiers
│
├── app/
│   ├── checker.py            # Daily expiry checker — iterates vehicles, sends Telegram/Email/Push
│   ├── scheduler.py          # APScheduler background scheduler (daily 09:00 cron + 9h auto-fetch)
│   ├── email_sender.py       # Gmail SMTP SSL sender (reminder, expired, monthly summary)
│   ├── email_templates.py    # HTML email templates for reminder notifications
│   ├── telegram_bot.py       # Telegram bot message sender with formatted alerts
│   ├── push_notifications.py # Firebase Admin SDK — per-vehicle and bulk push notifications
│   ├── reminder_routes.py    # REST endpoints: send_reminder, bulk_reminders, notification_logs
│   ├── auto_fetcher.py       # Background auto-fetch of vehicle data from Parivahan
│   ├── notifier.py           # Legacy telegram sender (unused, kept for reference)
│   └── main.py               # Legacy entry point (unused, kept for reference)
│
├── routes/
│   ├── auth.py               # Authentication: login, register, forgot/reset password
│   ├── vehicles.py           # Vehicle CRUD, Parivahan fetch trigger, delete/restore
│   ├── users.py              # User management, role updates, profile updates
│   ├── subscriptions.py      # Razorpay order creation and payment verification
│   └── exports.py            # Excel export for users and vehicles
│
├── templates/                # Legacy Jinja2 HTML templates (login, index, edit)
├── uploads/                  # Vehicle document uploads
├── data/                     # Log files (push_logs.txt, email_logs.txt, telegram_logs.txt)
│
└── mv-tax-frontend/          # React 19 + Vite 8 frontend
    ├── .env                  # Firebase config vars (VITE_APP_*) and Razorpay key
    ├── index.html            # HTML entry with dark-theme inline script + style
    ├── vite.config.js        # Vite configuration
    ├── tailwind.config.js    # Tailwind with darkMode: 'class'
    ├── postcss.config.js     # PostCSS with Tailwind plugin
    ├── public/               # Static assets (favicon, vite.svg, firebase-messaging-sw.js)
    ├── src/
    │   ├── App.jsx           # React Router with React.lazy + Suspense code-splitting
    │   ├── main.jsx          # Entry — renders App, calls setupNotifications()
    │   ├── ThemeContext.jsx   # React context for dark/light theme, persisted to localStorage
    │   ├── notificationService.js  # FCM token acquisition + foreground message handler
    │   ├── firebase.js       # Firebase app init, getMessaging, getToken, onMessage
    │   ├── index.css         # Tailwind directives + custom keyframe animations
    │   ├── pages/
    │   │   ├── Dashboard.jsx       # Stats cards, PieChart, BarChart, health score
    │   │   ├── Vehicles.jsx        # Vehicle list with search, status filter, pagination
    │   │   ├── Users.jsx           # User CRUD with role/subscription management
    │   │   ├── Settings.jsx        # Profile update (username, email, phone, password)
    │   │   ├── Plans.jsx           # Subscription plan cards with Razorpay checkout
    │   │   ├── Notifications.jsx   # Notification log viewer (from notification_logs table)
    │   │   ├── Login.jsx, Register.jsx, ForgotPassword.jsx, ResetPassword.jsx
    │   │   ├── DeletedVehicles.jsx, DeletedUsers.jsx
    │   └── components/
    │       ├── Sidebar.jsx         # Collapsible sidebar + mobile hamburger overlay
    │       ├── StatsCard.jsx       # Memoized gradient stat card
    │       ├── VehicleTable.jsx    # Vehicle rows with Fetch/Mail/Push/Edit/History/WhatsApp
    │       ├── AddVehicleForm.jsx  # Add vehicle form with email field
    │       ├── EditVehicleModal.jsx
    │       ├── FetchModal.jsx      # Animated terminal-log modal for Parivahan fetch
    │       ├── HistoryModal.jsx    # Vehicle edit history viewer
    │       └── ProtectedRoute.jsx  # Auth guard + role-based access control"""

for line in structure.split('\n'):
    doc.add_paragraph(line, style='Normal')

doc.add_paragraph('')

# ============================================================
# 3. DATABASE SCHEMA
# ============================================================
doc.add_heading('3. Database Schema', level=1)
doc.add_paragraph('SQLite database (vehicles.db) with 7 tables and WAL mode for concurrent access.')

doc.add_heading('3.1 Tables', level=2)

tables = [
    ('users', '''CREATE TABLE IF NOT EXISTS users (
    id INTEGER PRIMARY KEY,
    username TEXT UNIQUE,
    email TEXT UNIQUE,
    phone TEXT UNIQUE,
    password TEXT,
    role TEXT,
    subscription_status TEXT DEFAULT 'free',
    subscription_expiry TEXT
)''', 'Stores all user accounts with role-based access (admin/staff/viewer) and subscription info. Default users: admin, staff1, viewer1.'),
    ('vehicles', '''CREATE TABLE IF NOT EXISTS vehicles (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    vehicle_number TEXT,
    expiry_date TEXT,
    phone TEXT,
    owner TEXT,
    chassis_last5 TEXT,
    state_name TEXT,
    support_document TEXT,
    vahan_owner_name TEXT,
    added_by TEXT,
    email TEXT,
    notification_enabled INTEGER DEFAULT 1
)''', 'Core vehicle registry. Each vehicle has tax expiry date, chassis (last 5 digits for Parivahan), Parivahan owner name fetched via scraper, and email for reminders.'),
    ('deleted_vehicles', 'Mirrors vehicles table plus deleted_by/deleted_at. Stores soft-deleted vehicles for admin restore.', ''),
    ('vehicle_history', 'Tracks edit history on vehicles: old values + edited_at timestamp per change.', ''),
    ('deleted_users', 'Mirrors users table. Stores soft-deleted users for admin restore.', ''),
    ('payments', '''CREATE TABLE IF NOT EXISTS payments (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_id INTEGER,
    razorpay_payment_id TEXT,
    razorpay_order_id TEXT,
    razorpay_signature TEXT,
    amount INTEGER,
    currency TEXT,
    plan TEXT,
    status TEXT,
    paid_at TEXT
)''', 'Records Razorpay payment transactions for subscription upgrades.'),
    ('fcm_tokens', 'Stores Firebase Cloud Messaging tokens per username for push notification delivery.', ''),
    ('notification_logs', 'Audit log of all sent notifications: vehicle_number, channel, message, success status, sent_at.', ''),
]
for name, schema, desc in tables:
    doc.add_heading(f'{name}', level=3)
    if schema:
        for line in schema.split('\n'):
            doc.add_paragraph(line, style='Normal')
    if desc:
        doc.add_paragraph(desc)

doc.add_heading('3.2 Database Indexes', level=2)
indexes = [
    'idx_users_username ON users(username)',
    'idx_users_phone ON users(phone)',
    'idx_vehicles_vehicle_number ON vehicles(vehicle_number)',
    'idx_vehicles_added_by ON vehicles(added_by)',
    'idx_vehicles_expiry_date ON vehicles(expiry_date)',
    'idx_vehicle_history_vehicle_id ON vehicle_history(vehicle_id)',
    'idx_payments_user_id ON payments(user_id)',
]
for idx in indexes:
    doc.add_paragraph(f'• CREATE INDEX IF NOT EXISTS {idx}', style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 4. API ENDPOINTS
# ============================================================
doc.add_heading('4. API Endpoints', level=1)

endpoints = [
    ('Auth', [
        ('POST /api/login', 'Authenticate user, returns JWT token (7-day expiry)'),
        ('POST /api/register', 'Register new user account'),
        ('POST /api/forgot_password', 'Send password reset link via email'),
        ('POST /api/reset_password', 'Reset password with token'),
    ]),
    ('Vehicles', [
        ('GET /api/vehicles', 'List all vehicles (filtered by role — admin sees all, others see own)'),
        ('POST /api/add_vehicle', 'Add new vehicle (multipart/form-data with optional support document)'),
        ('PUT /api/update_vehicle/<id>', 'Update vehicle details (owner, phone, email, chassis, expiry)'),
        ('DELETE /api/delete_vehicle/<id>', 'Soft-delete vehicle (moves to deleted_vehicles)'),
        ('GET /api/fetch_vehicle_info/<number>', 'Trigger Parivahan scraper via subprocess, returns JSON with tax_upto'),
        ('GET /api/vehicle_history/<id>', 'Get edit history for a specific vehicle'),
        ('GET /api/deleted_vehicles', 'List soft-deleted vehicles (admin only)'),
        ('POST /api/restore_vehicle/<id>', 'Restore soft-deleted vehicle'),
        ('DELETE /api/permanent_delete_vehicle/<id>', 'Permanently delete vehicle'),
    ]),
    ('Users', [
        ('GET /api/users', 'List all users (admin only)'),
        ('POST /api/add_user', 'Create new user with role (admin only)'),
        ('PUT /api/update_user/<id>', 'Update user profile (admin only)'),
        ('DELETE /api/delete_user/<id>', 'Soft-delete user (admin only)'),
        ('POST /api/restore_user/<id>', 'Restore soft-deleted user'),
        ('GET /api/deleted_users', 'List soft-deleted users'),
        ('PUT /api/update_user_role/<id>', 'Change user role (staff/viewer)'),
        ('PUT /api/update_profile', 'Update own profile (username, email, phone, password)'),
        ('POST /api/user_subscription/<id>', 'Upgrade/extend/downgrade user subscription (admin only)'),
    ]),
    ('Dashboard', [
        ('GET /api/dashboard_stats', 'Aggregated stats: vehicle counts by status, user analytics, health score'),
        ('GET /api/reminder_stats', 'Reminder readiness stats: vehicles needing reminders, with email, notifications on'),
    ]),
    ('Notifications', [
        ('POST /api/send_reminder', 'Send reminder for a vehicle (channel: telegram/email/push/all)'),
        ('POST /api/send_bulk_reminders', 'Send bulk reminders to all vehicles (admin only)'),
        ('POST /api/save_fcm_token', 'Save browser FCM push token for current user'),
        ('GET /api/notification_logs', 'Get last 100 notification log entries'),
    ]),
    ('Subscriptions', [
        ('POST /api/create_order', 'Create Razorpay order for Staff Pro subscription'),
        ('POST /api/verify_payment', 'Verify Razorpay payment signature and activate subscription'),
    ]),
    ('Exports', [
        ('GET /api/export_users', 'Export all users to Excel (.xlsx)'),
        ('GET /api/export_user/<id>', 'Export single user\'s vehicles to Excel'),
    ]),
]

for group, routes in endpoints:
    doc.add_heading(group, level=2)
    for route, desc in routes:
        p = doc.add_paragraph()
        p.add_run(f'{route}: ').bold = True
        p.add_run(desc)

doc.add_paragraph('')

# ============================================================
# 5. PARIVAHAN SCRAPER
# ============================================================
doc.add_heading('5. Parivahan Scraper (fetch_vehicle.py)', level=1)
doc.add_paragraph(
    'The scraper uses Playwright (Chromium) to automate the Indian government Parivahan portal '
    '(parivahan.gov.in) for fetching vehicle tax expiry dates. It is invoked as a subprocess '
    'by the Flask backend at GET /api/fetch_vehicle_info/<vehicle_number>.'
)

doc.add_heading('5.1 Workflow', level=2)
steps = [
    'Parse state code from vehicle number prefix (e.g., MH → Maharashtra). Supports 37 states/UTs.',
    'Launch headless Chromium with anti-detection flags (incognito, custom user-agent, viewport).',
    'Navigate to parivahan.gov.in vehicle services page with up to 3 retry attempts.',
    'Select state from dropdown and wait for VAHAN sub-page to load.',
    'Handle "Previous session active" popup if present.',
    'Close informational popup if present.',
    'Click "Vehicle Registration No." with 3 retry attempts and screenshots on failure.',
    'Enter vehicle number, check consent checkbox, click Proceed.',
    'Detect CAPTCHA — if present, screenshot and exit.',
    'Handle authentication dialog (find visible "Proceed" button in dialog overlays).',
    'Click "Pay Your Tax" trigger button.',
    'Enter chassis last 5 digits and click "Verify Details".',
    'Await result page — check for challan dialog.',
    'If challan detected: press Escape key → click backdrop (10,10) → JavaScript remove .ui-dialog. If all fail, return challan_pending: true.',
    'If no challan or dismissed: click #form_eapp\\:tf_show_button submit button.',
    'Extract #taxFrom\\:tf_tax_upto (expiry date) and #taxFrom\\:tf_owner_name (owner name).',
    'Output JSON to stdout for parent process to parse.',
    'On any error: screenshot, print traceback, ensure browser.close().',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('5.2 Error Handling', level=2)
errors = [
    'Network failure: 3 retries with 5s backoff, then raise Exception.',
    'Element timeout: each wait_for_selector has 15-30s timeout.',
    'Vehicle Registration click: 3 retries with full-page screenshot per attempt.',
    'CAPTCHA: screenshot saved as captcha.png, clean exit.',
    'Challan: 3-step dismiss (Escape → click backdrop → JS remove), fallback JSON output with challan_pending: true.',
    'Generic exception: traceback.print_exc(), error screenshot, browser closed.',
    'Subprocess timeout: Flask route has 180s timeout — returns HTTP 408.',
]
for err in errors:
    doc.add_paragraph(f'• {err}', style='List Bullet')

doc.add_heading('5.3 State Support (37 states/UTs)', level=2)
states = ['WB', 'BR', 'OD', 'PB', 'HR', 'KL', 'TS', 'JK', 'UK', 'GA', 'MH', 'MP', 'DL', 'UP', 'RJ', 'CG', 'GJ', 'KA', 'TN', 'AP', 'NL', 'AS', 'HP', 'PY', 'AN', 'CH', 'DN', 'DD', 'LD', 'MN', 'ML', 'MZ', 'OR', 'SK', 'TR', 'AR', 'JH']
doc.add_paragraph(', '.join(states))

doc.add_paragraph('')

# ============================================================
# 6. NOTIFICATION SYSTEM
# ============================================================
doc.add_heading('6. Notification System', level=1)

doc.add_heading('6.1 Channel Details', level=2)

channels = [
    ('Gmail SMTP (app/email_sender.py)', 
     'Connects via smtplib.SMTP_SSL to smtp.gmail.com:465. Supports send_reminder_email, '
     'send_expired_email, and send_monthly_summary. HTML templates in email_templates.py. '
     'Uses EMAIL_ADDRESS and EMAIL_PASSWORD from .env (Gmail App Password).'),
    ('Telegram Bot (app/telegram_bot.py)',
     'Sends formatted messages via Telegram Bot API (https://api.telegram.org/bot<token>/sendMessage). '
     'Uses TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID from .env. Messages include urgency emojis '
     '(red circle for expired, orange for urgent, green for warning).'),
    ('Push (app/push_notifications.py)',
     'Uses Firebase Admin SDK with service account authentication (firebase-service-account.json). '
     'Per-vehicle push via send_vehicle_push(token, vehicle_number, days_left) and bulk via '
     'send_bulk_push(tokens, title, body). Frontend registers FCM tokens via /api/save_fcm_token. '
     'VAPID key from Firebase Web Push certificates configured in frontend .env.'),
    ('WhatsApp (helpers.py)',
     'Uses UltraMSG API (api.ultramsg.com). send_whatsapp_message(phone, message) sends messages. '
     'Currently not wired into cron or API routes. Frontend provides wa.me deep-link button as '
     'client-side alternative.'),
]

for name, desc in channels:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

doc.add_heading('6.2 Reminder Tiers', level=2)
doc.add_paragraph(
    'The system uses configurable reminder tiers via ALERT_DAYS (default 7) and REMINDER_DAYS [7, 3, 1]:'
)
tiers = [
    ('safe', 'Days left > 7 — no reminder sent'),
    ('warning', 'Days left <= 7 — early warning'),
    ('urgent', 'Days left <= 3 — urgent notification'),
    ('final', 'Days left <= 1 — final reminder'),
    ('expired', 'Days left < 0 — expired notification'),
]
for tier, desc in tiers:
    doc.add_paragraph(f'• {tier}: {desc}', style='List Bullet')

doc.add_heading('6.3 Daily Cron (APScheduler)', level=2)
doc.add_paragraph(
    'The APScheduler runs in a daemon thread started from dashboard.py on app boot. Schedule:'
)
cron_items = [
    'Daily at SCHEDULE_TIME (default 09:00): check_expiry() — iterates all vehicles, sends Telegram + Email + Push reminders for non-safe vehicles.',
    'Every 9 hours: auto_fetch_vehicles() — re-fetches tax data from Parivahan for vehicles due within 7 days.',
    'On startup: runs check_expiry() once immediately + auto_downgrade_expired().',
]
for item in cron_items:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 7. SUBSCRIPTION & BILLING
# ============================================================
doc.add_heading('7. Subscription & Billing', level=1)
doc.add_paragraph(
    'Two-tier system: Free (viewer role) and Staff Pro at ₹19/month (staff role). '
    'Payments handled via Razorpay Web SDK. The Plans page displays both tiers with feature comparisons.'
)

doc.add_heading('7.1 Feature Comparison', level=2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Feature', 'Free (Viewer)', 'Staff Pro (₹19/mo)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
features = [
    ('View Vehicles', 'Yes', 'Yes'),
    ('Add Vehicles', 'No', 'Yes'),
    ('Edit Vehicles', 'No', 'Yes'),
    ('Delete Vehicles', 'No', 'Yes'),
    ('Email Reminder', 'No', 'Yes'),
    ('Export Excel', 'No', 'Yes'),
    ('WhatsApp Reminder', 'No', 'Yes'),
]
for r, (feat, free, pro) in enumerate(features, 1):
    table.rows[r].cells[0].text = feat
    table.rows[r].cells[1].text = free
    table.rows[r].cells[2].text = pro

doc.add_paragraph('')
doc.add_heading('7.2 Payment Flow', level=2)
payment_steps = [
    'Frontend calls POST /api/create_order → Flask creates Razorpay order.',
    'Razorpay checkout modal opens with order ID, amount (₹19 = 1900 paise).',
    'User completes payment → Razorpay calls handler with payment_id, order_id, signature.',
    'Frontend POSTs to /api/verify_payment → Flask verifies HMAC SHA256 signature.',
    'On success: user role updated to staff, subscription_status = active, expiry = 30 days.',
    'localStorage updated on frontend, page reloads to reflect new role.',
]
for step in payment_steps:
    doc.add_paragraph(f'• {step}', style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 8. FRONTEND COMPONENTS
# ============================================================
doc.add_heading('8. Frontend Architecture', level=1)

doc.add_heading('8.1 Routing & Code-Splitting (App.jsx)', level=2)
doc.add_paragraph(
    'All 11 page components are lazy-loaded via React.lazy() with a Suspense fallback '
    '(animated "Loading..." spinner). This splits the production bundle from a single 927 kB file '
    'into 18 optimized chunks:'
)
chunks = [
    'index-CAvgGY5m.js (472 kB) — shared deps: React, Router, Axios, Firebase, lucide-react',
    'Dashboard-DmRKCuTw.js (383 kB) — recharts charting library',
    'Vehicles-auHUZvVx.js (27 kB) — vehicle management',
    'Users-VrwGXL8t.js (14 kB) — user management',
    'Register-DYmFlVl_.js (9 kB), Plans-DrBs-zf3.js (8 kB), Login-B-HfKhHl.js (7 kB)',
    'Other pages: 1-4 kB each',
    'Total gzip: ~154 kB',
]
for chunk in chunks:
    doc.add_paragraph(f'• {chunk}', style='List Bullet')

doc.add_heading('8.2 Key Components', level=2)
components = [
    ('Sidebar', 'Collapsible sidebar (w-64 / w-20) with framer-motion AnimatePresence. '
                'Mobile (< 1024px): hamburger overlay with slide-in + backdrop. '
                'Nav items: Dashboard, Vehicles, Users (admin), Deleted Vehicles, Deleted Users (admin), '
                'Notifications, Plans, Settings. Theme toggle (Sun/Moon) at bottom. '
                'Wrapped in React.memo for render optimization.'),
    ('Dashboard', 'Animated stat cards with gradient backgrounds (Total Vehicles, Expired, Expiring, Active). '
                  'Admin-only Users Overview section. Recharts PieChart (donut) + BarChart for vehicle status. '
                  'User Analytics bar chart + subscription PieChart. Health score calculation. '
                  'Uses useMemo for chart data, useCallback for fetchStats.'),
    ('VehicleTable', 'Full vehicle table with status badges (color-coded: green/yellow/red). '
                     'Action buttons per row: Fetch (Parivahan scraper), Mail (email reminder), '
                     'Push (browser push), Edit, History, WhatsApp (wa.me deep-link), Delete (admin). '
                     'Uses FetchModal for animated terminal-log display.'),
    ('Plans', 'Two pricing cards (Free / Staff Pro) with framer-motion stagger animation. '
              'Active subscription banner showing remaining days. '
              'Razorpay checkout integration.'),
    ('Notifications', 'Fetches from GET /api/notification_logs. Displays recent reminders with channel '
                      'icons (Mail, MessageSquare, Smartphone), success/fail status, timestamps.'),
]
for name, desc in components:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

doc.add_heading('8.3 Performance Optimizations', level=2)
perf_items = [
    'React.lazy + Suspense code-splitting (18 production chunks)',
    'React.memo on StatsCard and Sidebar to prevent unnecessary re-renders',
    'useMemo for chart data, filtered vehicle lists, paginated results',
    'useCallback for all async fetch functions (fetchVehicles, fetchStats)',
    'Arrow functions moved out of JSX render (deleted vehicles/users modals)',
    'All useEffect/useMemo/useCallback have complete dependency arrays',
    'Unused imports removed (ArrowLeft from Plans, React from main.jsx)',
    'window.location.reload() replaced with delayed modal close',
    '21 console.log statements removed from all frontend files',
]
for item in perf_items:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 9. DARK/LIGHT THEME
# ============================================================
doc.add_heading('9. Dark/Light Theme', level=1)
doc.add_paragraph(
    'Tailwind CSS darkMode: "class" strategy. The theme is toggled via a Sun/Moon button in the Sidebar '
    'and persisted to localStorage under the "theme" key.'
)
theme_items = [
    '<html class="dark"> hardcoded in index.html as default',
    'Inline <script> in index.html: removes dark class only if localStorage says "light"',
    'Inline <style> in index.html: sets background-color: black immediately to prevent flash',
    'ThemeContext.jsx: React context providing theme + toggleTheme, toggles class on document.documentElement',
    'All 14+ page and component files use dark: Tailwind variants',
    'Dark backgrounds: bg-gray-950, containers: dark:bg-gray-900/80, inputs: dark:bg-gray-800',
]
for item in theme_items:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 10. ENVIRONMENT CONFIGURATION
# ============================================================
doc.add_heading('10. Environment Configuration', level=1)

doc.add_heading('10.1 Backend (.env)', level=2)
backend_env = [
    ('FLASK_SECRET_KEY', 'Flask session secret key (auto-generated if not set)'),
    ('JWT_SECRET_KEY', 'JWT signing key (auto-generated if not set)'),
    ('CORS_ORIGINS', 'Comma-separated allowed origins for CORS'),
    ('TELEGRAM_BOT_TOKEN', 'Telegram Bot API token from @BotFather'),
    ('TELEGRAM_CHAT_ID', 'Target Telegram chat/group ID for notifications'),
    ('EMAIL_ADDRESS', 'Gmail address for SMTP sending'),
    ('EMAIL_PASSWORD', 'Gmail App Password (16-char, no spaces)'),
    ('FIREBASE_SERVICE_ACCOUNT_PATH', 'Path to Firebase Admin SDK service account JSON'),
    ('SCHEDULE_TIME', 'Daily cron time (default: 09:00)'),
]
for var, desc in backend_env:
    p = doc.add_paragraph()
    p.add_run(f'{var}: ').bold = True
    p.add_run(desc)

doc.add_heading('10.2 Frontend (.env)', level=2)
frontend_env = [
    ('VITE_FIREBASE_VAPID_KEY', 'VAPID key from Firebase Cloud Messaging Web Push certs'),
    ('VITE_FIREBASE_API_KEY', 'Firebase Web API key'),
    ('VITE_FIREBASE_AUTH_DOMAIN', 'Firebase auth domain'),
    ('VITE_FIREBASE_PROJECT_ID', 'Firebase project ID'),
    ('VITE_FIREBASE_STORAGE_BUCKET', 'Firebase storage bucket'),
    ('VITE_FIREBASE_MESSAGING_SENDER_ID', 'Firebase sender ID'),
    ('VITE_FIREBASE_APP_ID', 'Firebase app ID'),
    ('VITE_FIREBASE_MEASUREMENT_ID', 'Google Analytics measurement ID'),
    ('VITE_RAZORPAY_KEY_ID', 'Razorpay API key (test/live)'),
]
for var, desc in frontend_env:
    p = doc.add_paragraph()
    p.add_run(f'{var}: ').bold = True
    p.add_run(desc)

doc.add_paragraph('')

# ============================================================
# 11. SECURITY
# ============================================================
doc.add_heading('11. Security Considerations', level=1)
security = [
    'Passwords hashed with bcrypt before storage',
    'JWT tokens expire after 7 days; axios interceptor redirects to /login on expiry',
    'All API routes (except login/register/forgot/reset) require JWT via @jwt_required()',
    'Role-based access control: admin-only routes check user role before processing',
    'ProtectedRoute component guards frontend routes — redirects to /login if no token',
    'CORS restricted to ALLOWED_ORIGINS from .env',
    'Flask secret key auto-generated via os.urandom(32) as fallback',
    'Firebase service account JSON excluded via .gitignore',
    '.env excluded via .gitignore to prevent secret leakage',
    'Razorpay payment signature verified server-side via HMAC SHA256',
]
for item in security:
    doc.add_paragraph(f'• {item}', style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 12. BUILD & DEPLOYMENT
# ============================================================
doc.add_heading('12. Build & Deployment', level=1)

doc.add_heading('12.1 Build', level=2)
build_steps = [
    'Frontend: cd mv-tax-frontend && npm run build → outputs to dist/',
    'Production build: 22 chunks, total ~927 kB JS → ~154 kB gzip',
    'Backend: python dashboard.py (or gunicorn dashboard:app for production)',
]
for step in build_steps:
    doc.add_paragraph(f'• {step}', style='List Bullet')

doc.add_heading('12.2 Dependencies', level=2)
p = doc.add_paragraph()
p.add_run('Python: ').bold = True
p.add_run('Flask, flask-jwt-extended, flask-cors, APScheduler, playwright, firebase-admin, '
          'bcrypt, python-dotenv, reportlab, openpyxl, requests, gunicorn')
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Node: ').bold = True
p.add_run('React 19, Vite 8, Tailwind CSS 3, framer-motion, recharts, lucide-react, '
          'axios, react-router-dom, react-hot-toast, firebase')

doc.add_paragraph('')
doc.add_paragraph('─' * 80)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('— End of Report —').italic = True

# Save
output_path = 'D:\\Vehicle Reminder\\MV_Tax_Reminder_Project_Report.docx'
doc.save(output_path)
print(f"Report saved to: {output_path}")
